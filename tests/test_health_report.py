"""Business Health Report (.docx) tests: generation + shareable anonymity."""
import pytest
from docx import Document

from src.report.health_report import FIRM_REAL, generate_health_report
from src.warehouse.db import DB_PATH

_skip = pytest.mark.skipif(not DB_PATH.exists(), reason="no warehouse db")


def _text(path):
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return " ".join(parts)


@_skip
def test_internal_report_generates_with_sections():
    path = generate_health_report("internal")
    d = Document(path)
    text = _text(path)
    assert "Business Health Report" in text
    assert "Executive summary" in text
    assert "Key risks" in text and "Key opportunities" in text
    assert "Prioritised recommendations" in text
    assert len(d.tables) >= 4         # trends, risks, opportunities, recs, 2x2
    assert FIRM_REAL in text          # real firm name in internal


@_skip
def test_shareable_report_is_anonymised():
    path = generate_health_report("shareable")
    text = _text(path)
    assert FIRM_REAL not in text                      # firm name withheld
    assert "anonymised" in text.lower()
    # no full mobile numbers in the body
    import re
    assert not re.search(r"(?<!\d)[6-9]\d{9}(?!\d)", text)
