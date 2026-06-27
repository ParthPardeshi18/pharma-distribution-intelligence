"""Master vision document — Platform Vision & Architecture v1.0 (.docx).

Renders architecture diagrams (matplotlib) and consolidates the platform vision,
architecture, GIS/AI roadmaps, ERP expansion, SaaS vision, and maturity model into
a portfolio-safe Word document. Firm-neutral (no real identities) so it is
committable as a portfolio piece.

Run:  python -m src.report.vision
"""
from __future__ import annotations

import datetime as dt

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrow, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.utils import PROJECT_ROOT

DIAGRAMS = PROJECT_ROOT / "docs" / "diagrams"
OUT = PROJECT_ROOT / "reports" / "shareable" / "Platform_Vision_and_Roadmap_v1.0.docx"
_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_GREY = RGBColor(0x55, 0x55, 0x55)


# ----------------------------- diagrams ----------------------------------- #
def _box(ax, x, y, w, h, text, fc="#dce6f2", ec="#1f3a5f", fs=9, bold=False):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02",
                                fc=fc, ec=ec, lw=1.2))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fs, fontweight="bold" if bold else "normal", wrap=True)


def _arrow(ax, x1, y1, x2, y2, color="#888"):
    ax.add_patch(FancyArrow(x1, y1, x2 - x1, y2 - y1, width=0.004,
                            head_width=0.02, head_length=0.02,
                            length_includes_head=True, color=color))


def render_diagrams() -> list[str]:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    made = []

    # 1. layered stack
    fig, ax = plt.subplots(figsize=(7.5, 9))
    layers = [
        ("Source systems — ERP exports · geo · FX", "#eef2f7"),
        ("1 · ERP Adapter Layer (one ERP boundary)", "#dce6f2"),
        ("2 · Data Ingestion — clean · resolve · keys", "#dce6f2"),
        ("3 · Data Warehouse — star schema · lineage · GeoJSON", "#cdddee"),
        ("4 · Business Intelligence — KPI registry · DQ", "#dce6f2"),
        ("5 · Decision Intelligence — risks · recs · Health Index", "#cfe8d6"),
        ("6 · Strategic Analytics — ABC/RFM · forecasting", "#cfe8d6"),
        ("7 · Geographic Intelligence — canonical GeoJSON · maps", "#cfe8d6"),
        ("8 · AI Layer (roadmap) — forecast · risk · copilot", "#f5e6c8"),
        ("9 · Presentation — Power BI · .docx · APIs · web", "#dce6f2"),
    ]
    n = len(layers)
    for i, (txt, fc) in enumerate(layers):
        y = 1 - (i + 1) * (1 / (n + 1))
        _box(ax, 0.08, y, 0.74, 1 / (n + 1) * 0.8, txt, fc=fc, fs=9,
             bold=i in (3, 5, 7))
        if i:
            _arrow(ax, 0.45, y + 1 / (n + 1) * 0.95, 0.45, y + 1 / (n + 1) * 0.8)
    _box(ax, 0.85, 0.30, 0.13, 0.40, "10 · Automation\n& Orchestration\n(run_pipeline)",
         fc="#f0d9d9", fs=8)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.set_title("Platform Architecture — layered stack", fontsize=13, color="#1f3a5f")
    p = DIAGRAMS / "architecture_stack.png"
    fig.tight_layout(); fig.savefig(p, dpi=120); plt.close(fig); made.append(str(p))

    # 2. data flow
    fig, ax = plt.subplots(figsize=(9, 3.2))
    flow = [("Raw\nexports", "#eef2f7"), ("Adapter\n(canonical)", "#dce6f2"),
            ("Warehouse\n(star)", "#cdddee"), ("DI · Strategic\n· GIS", "#cfe8d6"),
            ("Power BI ·\n.docx · maps", "#dce6f2")]
    w = 0.16
    for i, (txt, fc) in enumerate(flow):
        x = 0.02 + i * 0.20
        _box(ax, x, 0.35, w, 0.34, txt, fc=fc, fs=9, bold=True)
        if i:
            _arrow(ax, x - 0.04, 0.52, x, 0.52)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.set_title("Data flow — reconciled to the rupee at the warehouse boundary",
                 fontsize=12, color="#1f3a5f")
    p = DIAGRAMS / "data_flow.png"
    fig.tight_layout(); fig.savefig(p, dpi=120); plt.close(fig); made.append(str(p))

    # 3. GIS architecture
    fig, ax = plt.subplots(figsize=(9, 3.4))
    _box(ax, 0.02, 0.55, 0.20, 0.30, "Geocoding cache\n(lat/lon)", fc="#eef2f7", fs=9)
    _box(ax, 0.28, 0.55, 0.22, 0.30, "Canonical GeoJSON\nlayers (data/geo)", fc="#cfe8d6", fs=9, bold=True)
    _box(ax, 0.56, 0.55, 0.20, 0.30, "dim_geo_feature\n(geometry=GeoJSON)", fc="#cdddee", fs=9, bold=True)
    _box(ax, 0.80, 0.45, 0.18, 0.48, "Power BI · Leaflet ·\nMapbox · Azure ·\ndeck.gl · web", fc="#dce6f2", fs=8)
    _box(ax, 0.28, 0.12, 0.22, 0.26, "FUTURE: fact_gps_ping\nfact_vehicle_route", fc="#f5e6c8", fs=8)
    _arrow(ax, 0.22, 0.70, 0.28, 0.70); _arrow(ax, 0.50, 0.70, 0.56, 0.70)
    _arrow(ax, 0.76, 0.70, 0.80, 0.70); _arrow(ax, 0.39, 0.38, 0.60, 0.55, color="#b06a00")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.set_title("GIS architecture — GeoJSON canonical, future-proofed", fontsize=12, color="#1f3a5f")
    p = DIAGRAMS / "gis_architecture.png"
    fig.tight_layout(); fig.savefig(p, dpi=120); plt.close(fig); made.append(str(p))
    return made


# ----------------------------- docx helpers ------------------------------- #
def _h(doc, t, lvl=1):
    h = doc.add_heading(t, lvl)
    for r in h.runs:
        r.font.color.rgb = _NAVY
    return h


def _p(doc, t, size=10, italic=False, color=None, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.font.italic = italic; r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


# ----------------------------- builder ------------------------------------ #
def build() -> str:
    imgs = render_diagrams()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)

    _p(doc, "Pharmaceutical Distribution Intelligence Platform", bold=True,
       color=_NAVY, size=24).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p(doc, "Vision, Architecture & 5-Year Roadmap — Version 1.0", italic=True,
       color=_GREY, size=14).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p(doc, f"{dt.date.today():%B %Y}", color=_GREY, size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Vision
    _h(doc, "1 · Platform Vision", 1)
    _p(doc, "This platform began as an internal ERP reporting solution for a "
            "mid-size pharmaceutical distributor and has evolved into a complete "
            "Pharmaceutical Distribution Intelligence Platform: a reproducible, "
            "PII-safe system that turns raw ERP exports into a reconciled warehouse, "
            "decision intelligence, strategic analytics, geographic intelligence, "
            "BI exports, and an executive report — from one command.")
    _p(doc, "The vision over five years is to grow from single-firm decision support "
            "(v1.0) into a predictive, AI-assisted platform (v2.0) and ultimately an "
            "autonomous, multi-tenant SaaS product (v3.0) for the pharmaceutical "
            "distribution sector — without re-architecting, because the foundational "
            "contracts are stable.")

    # 2. Architecture
    _h(doc, "2 · Product Architecture", 1)
    _p(doc, "Ten loosely-coupled layers; each depends only on the contract below it, "
            "so the ERP, warehouse, BI tool, or AI models can change independently.")
    doc.add_picture(imgs[0], width=Inches(5.4))
    doc.add_picture(imgs[1], width=Inches(6.2))
    _table(doc, ["Layer", "Responsibility (v1.0)"], [
        ["1 ERP Adapter", "MediVision exports → canonical frames (one ERP boundary)"],
        ["2 Ingestion", "clean · reconcile · entity resolution · keys"],
        ["3 Warehouse", "SQLite star schema · lineage/audit/quality · spatial table"],
        ["4 Business Intelligence", "KPI registry · validation · DQ dashboard"],
        ["5 Decision Intelligence", "insights · risks · opportunities · recs · Health Index"],
        ["6 Strategic Analytics", "ABC/RFM · lifecycle · seasonality · forecasting"],
        ["7 Geographic Intelligence", "geocoding · canonical GeoJSON · maps"],
        ["8 AI Layer (roadmap)", "forecasting · risk · pricing · optimisation · copilot"],
        ["9 Presentation", "Power BI exports · .docx reports · (future) APIs/web"],
        ["10 Automation", "one-command pipeline → cron/CI → cloud"],
    ])

    # 3. GIS roadmap
    _h(doc, "3 · Future GIS Roadmap", 1)
    _p(doc, "v1.0 made GeoJSON the canonical spatial format with a warehouse spatial "
            "table and reserved GPS/vehicle-route tables. Enhancements add layers and "
            "modules, never a redesign.")
    doc.add_picture(imgs[2], width=Inches(6.2))
    _table(doc, ["Enhancement", "Priority"], [
        ["Official admin boundaries (GADM/OSM/Census)", "P1"],
        ["GPS-enabled delivery tracking", "P2"],
        ["Vehicle routing (actual paths)", "P2"],
        ["Drive-time analysis / isochrones", "P2"],
        ["Territory balancing", "P2"],
        ["Market penetration & white-space maps", "P2"],
        ["Route optimization (VRP)", "P3"],
        ["Warehouse location optimization", "P3"],
        ["Spatial AI (demand surfaces, hot-spots)", "P3"],
    ])

    # 4. AI roadmap
    _h(doc, "4 · AI Roadmap", 1)
    _p(doc, "Each AI module reads the warehouse + KPI registry and writes predictions "
            "back as new fact tables — no redesign. Every model ships with confidence, "
            "assumptions, quality metrics, and a business explanation.")
    _table(doc, ["AI module", "Value", "Complexity", "Priority"], [
        ["Sales forecasting", "cash/inventory/capacity planning", "Low", "P1"],
        ["Cash-flow forecasting", "30/60/90-day cash visibility", "Medium", "P1"],
        ["Customer risk prediction", "churn & credit risk → win-back", "Medium", "P1"],
        ["Demand forecasting", "SKU reorder & scheme planning", "High", "P2"],
        ["Inventory optimization", "min working capital at service level", "High", "P2"],
        ["Purchase recommendation", "auto reorder proposals", "Medium", "P2"],
        ["Pricing optimization", "margin lift on key SKUs", "High", "P2"],
        ["Supplier risk prediction", "continuity/price-risk flags", "Medium", "P2"],
        ["Territory recommendation", "route rebalancing & expansion", "Medium", "P2"],
        ["NL business assistant", "plain-language governed answers", "Medium", "P2"],
        ["Route optimization", "cut km/fuel/time (with GPS)", "High", "P3"],
        ["Executive Copilot", "proactive briefings & next actions", "High", "P3"],
    ])

    # 5. ERP expansion
    _h(doc, "5 · ERP Expansion", 1)
    _p(doc, "Adding an ERP means one adapter class emitting the canonical contract; "
            "the warehouse and everything above are untouched.")
    _table(doc, ["ERP", "Adapter effort"], [
        ["MediVision ✅ (v1.0)", "implemented"], ["Marg", "Low"], ["Busy", "Low"],
        ["Tally (XML/ODBC)", "Medium"], ["SAP (OData/RFC)", "High"],
        ["Oracle (SQL/REST)", "High"], ["Generic CSV", "Low (config-only)"]])

    # 6. SaaS vision
    _h(doc, "6 · SaaS Vision", 1)
    _p(doc, "The same codebase becomes multi-tenant by adding infrastructure and "
            "tenancy — not by rewriting the intelligence. The hard parts (canonical "
            "model, INR source-of-truth, anonymisation, GeoJSON, configurable "
            "analytics) already exist.")
    _table(doc, ["Capability", "SaaS design"], [
        ["Multi-company", "tenant_id + row-level security + per-tenant config"],
        ["Multi-branch", "branch dimension + consolidation views"],
        ["Multi-warehouse", "warehouses already modelled as geo features"],
        ["Multi-currency", "per-tenant base; INR-style reporting layer ✅"],
        ["Multi-language", "i18n string catalogs; localised KPI/insight templates"],
        ["Role-based security", "RBAC + role-based PII masking (anonymiser exists)"],
        ["Cloud deployment", "containers + Postgres + object storage + IaC/CI"],
        ["REST APIs", "/kpis /health-index /forecasts /geojson /reports"]])

    # 9. Maturity
    _h(doc, "7 · Maturity Assessment & Versions", 1)
    _p(doc, "v1.0 is a mature single-firm decision-support platform with a low-risk "
            "path to AI and SaaS.")
    _table(doc, ["Version", "Theme", "Introduces"], [
        ["v1.0 (now)", "Decision support", "Reconciled warehouse, DI + Health Index, "
         "strategic + forecasting, canonical GIS, Power BI/.docx, PII-safe, tested"],
        ["v2.0", "Predictive & deeper", "Line-item depth, AI (demand/cash/risk/pricing/"
         "inventory), official boundaries, NL assistant, Postgres + first APIs"],
        ["v3.0", "Autonomous & multi-tenant", "Route optimisation (GPS), Executive "
         "Copilot, multi-tenant SaaS (RBAC/cloud/i18n), spatial AI"]])
    _p(doc, "Guiding principle: each version adds layers and tables, never a "
            "redesign. This is the Version 1.0 release of a long-term platform — not "
            "the end of a project.", italic=True, color=_GREY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return str(OUT)


def main():
    print("Vision document:", build())


if __name__ == "__main__":
    main()
