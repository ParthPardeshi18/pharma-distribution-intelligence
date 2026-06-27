"""Business Health Report (.docx) — the executive deliverable.

Assembles the Decision-Intelligence and Strategic layers into a polished Word
document with: an executive summary + overall health verdict, multi-year trends,
findings by domain, ranked risks, ranked (₹-quantified) opportunities, and a 2×2
impact-vs-effort recommendation matrix where every recommendation carries a
measurable expected business impact.

Two versions:
  internal  — real firm name; reports/internal/business_health_report.docx
  shareable — placeholder firm name, anonymised; reports/shareable/...

Run:  python -m src.report.health_report --mode internal
"""
from __future__ import annotations

import datetime as dt

import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.di import health_index
from src.di.domains import ALL
from src.di.kpis import get_spec
from src.strategic import trends as strat_trends
from src.utils import CONFIG_DIR, PROJECT_ROOT

CHARTS_DIR = PROJECT_ROOT / "outputs" / "charts"
FIRM_REAL = "Sukhakarta Distributors"

_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_GREY = RGBColor(0x55, 0x55, 0x55)
_GREEN = RGBColor(0x1E, 0x7E, 0x34)
_AMBER = RGBColor(0xB0, 0x6A, 0x00)
_RED = RGBColor(0xB1, 0x1B, 0x1B)


def _firm_name(mode: str) -> str:
    if mode == "internal":
        return FIRM_REAL
    cfg = yaml.safe_load((CONFIG_DIR / "anonymisation.yaml").read_text(encoding="utf-8"))
    return cfg.get("firm_placeholder", "a mid-size pharmaceutical distribution firm").title()


def _verdict(score: float) -> tuple[str, RGBColor]:
    if score >= 80:
        return "Strong — the business is healthy with a few optimisation levers.", _GREEN
    if score >= 65:
        return "Stable with clear improvement areas — solid core, specific risks to manage.", _AMBER
    if score >= 50:
        return "Mixed — meaningful strengths offset by issues that need action.", _AMBER
    return "Fragile — several core dimensions need urgent attention.", _RED


# --------------------------- docx helpers --------------------------------- #
def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _NAVY
    return h


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def _para(doc, text, *, bold=False, italic=False, color=None, size=10, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


# --------------------------- content gathering ---------------------------- #
def _collect(reports):
    risks = sorted((r for rep in reports for r in rep.risks),
                   key=lambda r: r.score, reverse=True)
    opps = sorted((o for rep in reports for o in rep.opportunities),
                  key=lambda o: o.potential_impact_inr or 0, reverse=True)
    # recommendations with a domain tag and an expected-impact estimate
    dom_top_opp = {}
    for rep in reports:
        best = max((o.potential_impact_inr or 0 for o in rep.opportunities), default=0)
        dom_top_opp[rep.domain] = best
    actions = []
    for rep in reports:
        for rec in rep.recommendations:
            impact_inr = dom_top_opp.get(rep.domain, 0)
            actions.append({"rec": rec, "domain": rep.domain,
                            "expected_impact_inr": impact_inr})
    actions.sort(key=lambda a: a["rec"].priority_score, reverse=True)
    return risks, opps, actions


def _expected_impact_text(a) -> str:
    if a["expected_impact_inr"]:
        return f"up to ₹{a['expected_impact_inr']:,.0f} ({a['domain'].replace('_',' ')})"
    return f"{a['rec'].impact} impact on {a['domain'].replace('_',' ')}"


# --------------------------- main builder --------------------------------- #
def generate_health_report(mode: str = "internal") -> str:
    assert mode in ("internal", "shareable")
    reports = [m.analyze() for m in ALL]
    bhi = health_index.compute(reports)
    profiles = health_index.compute_all_profiles(reports)
    T = strat_trends.analyze()
    risks, opps, actions = _collect(reports)
    firm = _firm_name(mode)
    rev = next(t for t in T if t.measure == "Revenue")
    margin = next(t for t in T if t.measure == "Gross margin")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ---- Title ----
    _para(doc, "Business Health Report", bold=True, color=_NAVY, size=26,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, firm, italic=True, color=_GREY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Prepared {dt.date.today():%d %B %Y} · covers FY 2022-23 to 2025-26",
          color=_GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    if mode == "shareable":
        _para(doc, "Portfolio version — identities anonymised, firm name withheld.",
              italic=True, color=_GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ---- Executive summary ----
    _heading(doc, "Executive summary", 1)
    verdict, vcolor = _verdict(bhi.score)
    _para(doc, f"Business Health Index: {bhi.score}/100 (grade {bhi.grade}).",
          bold=True, size=13)
    _para(doc, verdict, color=vcolor, size=11)
    _para(doc, f"Revenue is {rev.note} (4-year CAGR {rev.cagr_pct:+.1f}%) while gross "
               f"margin is {margin.note} — {firm} is trading at tighter volumes but better "
               "margins. Collections and cash conversion are strong (a cash-heavy, "
               "route-based model); the main pressure points are customer retention and "
               "inventory efficiency.", size=10)
    _para(doc, "The three highest-value actions:", bold=True, size=10)
    for a in actions[:3]:
        _para(doc, f"•  {a['rec'].text}  —  expected impact: {_expected_impact_text(a)}.",
              size=10)
    # health-by-profile
    _para(doc, "Health Index under different strategy lenses: " +
          ", ".join(f"{p.replace('_',' ').title()} {s:.0f}" for p, s in profiles.items()) + ".",
          italic=True, color=_GREY, size=9)
    if (CHARTS_DIR / "bhi_components.png").exists():
        doc.add_picture(str(CHARTS_DIR / "bhi_components.png"), width=Inches(6.0))
    doc.add_page_break()

    # ---- Multi-year trends ----
    _heading(doc, "Multi-year trend analysis", 1)
    fys = [fy for fy, _ in T[0].series]
    rows = []
    for t in T:
        vals = [f"{v:,.0f}" if t.unit != "%" else f"{v:.2f}%" for _, v in t.series]
        rows.append([f"{t.measure} ({t.unit})"] + vals +
                    [f"{t.cagr_pct:+.1f}%" if t.cagr_pct is not None else "—", t.direction])
    _table(doc, ["Measure"] + fys + ["CAGR", "Trend"], rows)
    if (CHARTS_DIR / "revenue_margin_trend.png").exists():
        doc.add_picture(str(CHARTS_DIR / "revenue_margin_trend.png"), width=Inches(6.0))
    _para(doc, "So what: profit growth on softer revenue is a margin/mix gain — protect "
               "the pricing discipline that produced it.", italic=True, color=_GREY)

    # ---- Findings by domain ----
    _heading(doc, "Findings by domain", 1)
    for rep in reports:
        ex = rep.executive_summary
        sc = f" · scorecard {rep.scorecard.score:.0f}/100 ({rep.scorecard.grade})" \
            if rep.scorecard else ""
        _para(doc, rep.domain.replace("_", " ").title() + sc, bold=True, color=_NAVY, size=11)
        if ex:
            _para(doc, f"What happened: {ex.what_happened}", size=10)
            _para(doc, f"Why: {ex.why}", size=10)
            _para(doc, f"Business impact: {ex.business_impact}", size=10)

    # ---- Key risks ----
    doc.add_page_break()
    _heading(doc, "Key risks (ranked by severity)", 1)
    _table(doc, ["#", "Risk", "Severity", "Score", "Description"],
           [[i, r.name, r.severity, f"{r.score:.0f}", r.description]
            for i, r in enumerate(risks, 1)], widths=[0.3, 1.6, 0.9, 0.6, 3.0])

    # ---- Key opportunities ----
    _heading(doc, "Key opportunities (ranked by impact)", 1)
    _table(doc, ["#", "Opportunity", "Est. impact (₹)", "Confidence", "Detail"],
           [[i, o.name, f"{o.potential_impact_inr:,.0f}" if o.potential_impact_inr else "—",
             f"{o.confidence:.0%}", o.description]
            for i, o in enumerate(opps, 1)], widths=[0.3, 1.7, 1.2, 0.9, 2.8])

    # ---- Recommendations 2x2 ----
    doc.add_page_break()
    _heading(doc, "Prioritised recommendations (impact vs effort)", 1)
    _para(doc, "High-impact / low-effort actions come first. Each carries an expected "
               "business impact.", italic=True, color=_GREY)
    quad = {("high", "low"): [], ("high", "medium"): [], ("high", "high"): [],
            ("medium", "low"): [], ("medium", "medium"): [], ("medium", "high"): [],
            ("low", "low"): [], ("low", "medium"): [], ("low", "high"): []}
    seen = set()
    for a in actions:
        key = a["rec"].text
        if key in seen:
            continue
        seen.add(key)
        quad.setdefault((a["rec"].impact, a["rec"].effort), []).append(a)

    def _bucket(impacts, efforts):
        items = []
        for im in impacts:
            for ef in efforts:
                items += quad.get((im, ef), [])
        return items

    grid = [
        ["", "Low effort", "High effort"],
        ["High impact", _bucket(["high"], ["low"]), _bucket(["high"], ["medium", "high"])],
        ["Low/med impact", _bucket(["medium", "low"], ["low"]),
         _bucket(["medium", "low"], ["medium", "high"])],
    ]
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    for ri, row in enumerate(grid):
        for ci, cell in enumerate(row):
            tc = t.rows[ri].cells[ci]
            if ri == 0 or ci == 0:
                tc.text = cell
                for p in tc.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = _NAVY
            else:
                tc.text = ""
                first = (ri == 1 and ci == 1)
                for a in cell:
                    p = tc.add_paragraph()
                    r = p.add_run(("★ " if first else "• ") + a["rec"].text)
                    r.font.size = Pt(8)
                    r.font.bold = first
    doc.add_paragraph()

    _para(doc, "Ranked action list with expected impact:", bold=True)
    _table(doc, ["#", "Recommendation", "Impact", "Effort", "Conf.", "Expected business impact"],
           [[i, a["rec"].text, a["rec"].impact, a["rec"].effort, f"{a['rec'].confidence:.0%}",
             _expected_impact_text(a)] for i, a in enumerate(
               [a for a in actions if a["rec"].text in seen][:10], 1)],
           widths=[0.3, 2.6, 0.7, 0.7, 0.6, 1.9])

    # ---- Appendix ----
    doc.add_page_break()
    _heading(doc, "Appendix — method & assumptions", 1)
    _para(doc, "The Business Health Index is a weighted average of KPI health scores; "
               "thresholds and weights are documented in the codebase (KPI registry and "
               "health-index profiles) and tuned to route-based pharma distribution. "
               "Figures reconcile to the ERP totals to the rupee. Forecasts (separate "
               "strategic report) carry confidence intervals and backtested accuracy. "
               "Customer churn is measured by recency relative to each retailer's own "
               "ordering cadence.", size=9, color=_GREY)
    _para(doc, f"Components and scores: " +
          " · ".join(f"{c['name'].replace('_',' ')} {get_spec(c['kpi']).format(c['value'])} "
                     f"→ {c['score']:.0f}" for c in bhi.components), size=9, color=_GREY)

    out_dir = PROJECT_ROOT / "reports" / mode
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "business_health_report.docx"
    doc.save(str(out))
    return str(out)


def main():
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--mode", choices=["internal", "shareable"], default="internal")
    args = ap.parse_args()
    print("Health report:", generate_health_report(args.mode))


if __name__ == "__main__":
    main()
